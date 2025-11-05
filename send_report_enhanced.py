# --- [ENHANCED] Import libraries for email and Word export
import os
import boto3
import pandas as pd
from dotenv import load_dotenv
from datetime import datetime, timedelta
import re
from collections import Counter
import smtplib
import io
import argparse
import json
import tempfile
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.mime.application import MIMEApplication

# --- [NEW] Import for Word document creation
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

# --- [UNCHANGED] Plotly imports
import plotly.express as px
import plotly.graph_objects as go

# --- [NEW] Configuration for sections and time periods
AVAILABLE_SECTIONS = {
    'overview': 'Overall Platform Statistics',
    'engagement': 'Platform Engagement & Activity Trends',
    'adoption': 'Organizational Adoption',
    'features': 'Feature Adoption Analysis',
    'askai': 'Deep Dive: AskAI Usage'
}

# --- [UNCHANGED] Step 1: Data Loading and Preparation ---

# --- A. Data Fetching Function ---
def get_data_from_dynamodb(table_name, aws_key, aws_secret, aws_region):
    """Scans a DynamoDB table and returns a DataFrame."""
    print(f"  -> Fetching data from '{table_name}'...")
    try:
        dynamodb = boto3.resource(
            'dynamodb',
            aws_access_key_id=aws_key,
            aws_secret_access_key=aws_secret,
            region_name=aws_region
        )
        table = dynamodb.Table(table_name)
        response = table.scan()
        items = response['Items']
        while 'LastEvaluatedKey' in response:
            response = table.scan(ExclusiveStartKey=response['LastEvaluatedKey'])
            items.extend(response['Items'])
        df = pd.DataFrame(items)
        print(f"  ✅ Success! Found {len(df)} rows in '{table_name}'.")
        return df
    except Exception as e:
        print(f"  ❌ ERROR fetching data from {table_name}: {e}")
        return pd.DataFrame()

# --- [NEW] Data filtering function for time periods
def filter_data_by_time_period(df, start_date, end_date, date_column='createdAt'):
    """Filter DataFrame by time period."""
    if df.empty or date_column not in df.columns:
        return df

    # Convert timestamps to datetime if they're numeric
    if df[date_column].dtype in ['int64', 'float64']:
        df[date_column] = pd.to_datetime(df[date_column], unit='s', errors='coerce')
    else:
        df[date_column] = pd.to_datetime(df[date_column], errors='coerce')

    # Filter by date range
    mask = (df[date_column] >= start_date) & (df[date_column] <= end_date)
    return df[mask].copy()

# --- [NEW] Function to create Word document
def create_word_document(figures_dict, selected_sections, total_accounts, pro_users, team_users,
                        usage_log_valid_time_df, ask_ai_df, output_path="oak_report.docx"):
    """Create a Word document with the selected sections."""
    doc = Document()

    # Add title
    title = doc.add_heading('KaiShing OAK Usage Dashboard', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add report generation date
    date_para = doc.add_paragraph(f'Report Generated: {datetime.now().strftime("%d %B %Y, %H:%M:%S HKT")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add overview section if selected
    if 'overview' in selected_sections:
        doc.add_heading('Overall Platform Statistics', level=1)

        # Create a table for the KPI cards
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Add headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = f'Total Active Accounts\n{total_accounts}'
        hdr_cells[1].text = f'Pro Tier Users\n{pro_users}'
        hdr_cells[2].text = f'Team Tier Users\n{team_users}'

        # Center align the text in cells
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add engagement section if selected
    if 'engagement' in selected_sections:
        doc.add_heading('Platform Engagement & Activity Trends', level=1)

        if 'fig_wau' in figures_dict and figures_dict['fig_wau'] and figures_dict['fig_wau'].data:
            doc.add_paragraph('Weekly Active Users (WAU) Trend:')
            # Save figure as image and add to document
            try:
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                    img_path = tmp_file.name
                img_bytes = figures_dict['fig_wau'].to_image(format="png", width=800, height=450, scale=1)
                with open(img_path, 'wb') as f:
                    f.write(img_bytes)
                doc.add_picture(img_path, width=Inches(6))
                os.remove(img_path)  # Clean up temp file
            except Exception as e:
                print(f"  ⚠️  Warning: Could not export fig_wau: {e}")
                doc.add_paragraph('[Chart could not be exported]')

        if 'fig_heatmap' in figures_dict and figures_dict['fig_heatmap'] and figures_dict['fig_heatmap'].data:
            doc.add_paragraph('User Activity Heatmap (by Day and Hour):')
            try:
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                    img_path = tmp_file.name
                img_bytes = figures_dict['fig_heatmap'].to_image(format="png", width=800, height=450, scale=1)
                with open(img_path, 'wb') as f:
                    f.write(img_bytes)
                doc.add_picture(img_path, width=Inches(6))
                os.remove(img_path)
            except Exception as e:
                print(f"  ⚠️  Warning: Could not export fig_heatmap: {e}")
                doc.add_paragraph('[Chart could not be exported]')

    # Add adoption section if selected
    if 'adoption' in selected_sections:
        doc.add_heading('Organizational Adoption', level=1)

        if 'fig_site_activity' in figures_dict and figures_dict['fig_site_activity'] and figures_dict['fig_site_activity'].data:
            doc.add_paragraph('Activity Distribution by Site Code:')
            try:
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                    img_path = tmp_file.name
                img_bytes = figures_dict['fig_site_activity'].to_image(format="png", width=800, height=450, scale=1)
                with open(img_path, 'wb') as f:
                    f.write(img_bytes)
                doc.add_picture(img_path, width=Inches(6))
                os.remove(img_path)
            except Exception as e:
                print(f"  ⚠️  Warning: Could not export fig_site_activity: {e}")
                doc.add_paragraph('[Chart could not be exported]')

    # Add features section if selected
    if 'features' in selected_sections:
        doc.add_heading('Feature Adoption Analysis', level=1)

        if 'fig_features' in figures_dict and figures_dict['fig_features'] and figures_dict['fig_features'].data:
            doc.add_paragraph('What Features Are Users Exploring?:')
            try:
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                    img_path = tmp_file.name
                img_bytes = figures_dict['fig_features'].to_image(format="png", width=800, height=450, scale=1)
                with open(img_path, 'wb') as f:
                    f.write(img_bytes)
                doc.add_picture(img_path, width=Inches(6))
                os.remove(img_path)
            except Exception as e:
                print(f"  ⚠️  Warning: Could not export fig_features: {e}")
                doc.add_paragraph('[Chart could not be exported]')

    # Add AskAI section if selected
    if 'askai' in selected_sections:
        doc.add_heading('Deep Dive: AskAI Usage', level=1)

        if 'fig_askai_sites' in figures_dict and figures_dict['fig_askai_sites'] and figures_dict['fig_askai_sites'].data:
            doc.add_paragraph('AskAI Pioneers: Adoption by Site:')
            try:
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                    img_path = tmp_file.name
                img_bytes = figures_dict['fig_askai_sites'].to_image(format="png", width=800, height=450, scale=1)
                with open(img_path, 'wb') as f:
                    f.write(img_bytes)
                doc.add_picture(img_path, width=Inches(6))
                os.remove(img_path)
            except Exception as e:
                print(f"  ⚠️  Warning: Could not export fig_askai_sites: {e}")
                doc.add_paragraph('[Chart could not be exported]')

        if 'fig_askai_keywords' in figures_dict and figures_dict['fig_askai_keywords'] and figures_dict['fig_askai_keywords'].data:
            doc.add_paragraph('What Are Users Asking? Top AskAI Keywords:')
            try:
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                    img_path = tmp_file.name
                img_bytes = figures_dict['fig_askai_keywords'].to_image(format="png", width=800, height=450, scale=1)
                with open(img_path, 'wb') as f:
                    f.write(img_bytes)
                doc.add_picture(img_path, width=Inches(6))
                os.remove(img_path)
            except Exception as e:
                print(f"  ⚠️  Warning: Could not export fig_askai_keywords: {e}")
                doc.add_paragraph('[Chart could not be exported]')

    # Save the document
    doc.save(output_path)
    print(f"  ✅ Word document saved as: {output_path}")

# --- [NEW] Function to create interactive HTML report
def create_interactive_html(usage_log_df, ask_ai_df, selected_sections, total_accounts, pro_users, team_users,
                            output_path="oak_report.html"):
    """Create a standalone interactive HTML report with date selectors and presets."""
    print(f"  -> Building interactive HTML at {output_path} ...")

    def to_js_epoch(dt_series):
        ts = pd.to_datetime(dt_series, errors='coerce')
        return (ts.view('int64') // 10**6).astype('Int64')

    usage_records = []
    if not usage_log_df.empty:
        temp = usage_log_df.copy()
        if 'createdAt' in temp.columns and not pd.api.types.is_datetime64_any_dtype(temp['createdAt']):
            numeric_timestamps = pd.to_numeric(temp['createdAt'], errors='coerce')
            temp['createdAt'] = pd.to_datetime(numeric_timestamps, unit='s', errors='coerce')
        temp['createdAt_HKT'] = temp['createdAt'].dt.tz_convert('Asia/Hong_Kong')
        temp['createdAt_ms'] = to_js_epoch(temp['createdAt_HKT'])
        cols_present = {
            'account': temp['account'] if 'account' in temp.columns else pd.Series([None] * len(temp)),
            'site_code': temp['site_code'] if 'site_code' in temp.columns else pd.Series(['Unknown'] * len(temp)),
            'usage_type': temp['usage_type'] if 'usage_type' in temp.columns else pd.Series(['unknown'] * len(temp)),
            'createdAt_ms': temp['createdAt_ms']
        }
        packed = pd.DataFrame(cols_present).dropna(subset=['createdAt_ms'])
        usage_records = packed.to_dict(orient='records')

    askai_records = []
    if ask_ai_df is not None and not ask_ai_df.empty:
        temp = ask_ai_df.copy()
        if 'createdAt' in temp.columns and not pd.api.types.is_datetime64_any_dtype(temp['createdAt']):
            numeric_timestamps = pd.to_numeric(temp['createdAt'], errors='coerce')
            temp['createdAt'] = pd.to_datetime(numeric_timestamps, unit='s', errors='coerce')
        temp['createdAt_HKT'] = temp['createdAt'].dt.tz_convert('Asia/Hong_Kong')
        temp['createdAt_ms'] = to_js_epoch(temp['createdAt_HKT'])
        cols_present = {
            'account': temp['account'] if 'account' in temp.columns else pd.Series([None] * len(temp)),
            'site_code': temp['site_code'] if 'site_code' in temp.columns else pd.Series(['Unknown'] * len(temp)),
            'question': temp['question'] if 'question' in temp.columns else pd.Series([''] * len(temp)),
            'createdAt_ms': temp['createdAt_ms']
        }
        packed = pd.DataFrame(cols_present).dropna(subset=['createdAt_ms'])
        askai_records = packed.to_dict(orient='records')

    html_template = """
<!DOCTYPE html>
<html lang=\"en\">
<head>
  <meta charset=\"UTF-8\" />
  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\" />
  <title>KaiShing OAK Usage Dashboard - Interactive</title>
  <script src=\"https://cdn.plot.ly/plotly-2.32.0.min.js\"></script>
  <style>
    body { font-family: Arial, sans-serif; color: #333; margin: 20px; }
    h1 { text-align: center; }
    .controls { display: flex; flex-wrap: wrap; gap: 12px; align-items: center; justify-content: center; margin: 16px 0 24px; }
    .controls label { font-weight: 600; }
    .btn-group button { margin: 0 4px; padding: 6px 10px; border: 1px solid #ccc; background: #f7f7f7; border-radius: 6px; cursor: pointer; }
    .btn-group button.active { background: #00529B; color: #fff; border-color: #00529B; }
    .kpis { display: grid; grid-template-columns: repeat(3, 1fr); gap: 12px; margin: 12px 0 24px; }
    .kpi { background: #f4f4f4; border-radius: 8px; padding: 16px; text-align: center; }
    .kpi .value { font-size: 2em; }
    .chart { max-width: 1000px; margin: 24px auto; }
    .section-title { color: #00529B; border-bottom: 2px solid #00529B; padding-bottom: 6px; margin-top: 28px; }
  </style>
</head>
<body>
  <h1>KaiShing OAK Usage Dashboard</h1>
  <p style=\"text-align:center;color:#666;\">Report Generated: __GENERATED_AT__</p>

  <div class=\"controls\">
    <label for=\"startDate\">Start:</label>
    <input type=\"date\" id=\"startDate\" />
    <label for=\"endDate\">End:</label>
    <input type=\"date\" id=\"endDate\" />
    <div class=\"btn-group\">
      <button data-range=\"7\">Past 7 days</button>
      <button data-range=\"15\">Past 15 days</button>
      <button data-range=\"30\" class=\"active\">Past 30 days</button>
      <button data-range=\"all\">All time</button>
    </div>
    <button id=\"applyBtn\">Apply</button>
  </div>

  <div class=\"kpis\">
    <div class=\"kpi\"><div class=\"value\" id=\"kpi-total\">{total_accounts}</div><div>Total Active Accounts</div></div>
    <div class=\"kpi\"><div class=\"value\" id=\"kpi-pro\">{pro_users}</div><div>Pro Tier Users</div></div>
    <div class=\"kpi\"><div class=\"value\" id=\"kpi-team\">{team_users}</div><div>Team Tier Users</div></div>
  </div>

  <h2 class=\"section-title\">Platform Engagement & Activity Trends</h2>
  <div id=\"fig_wau\" class=\"chart\"></div>
  <div id=\"fig_heatmap\" class=\"chart\"></div>

  <h2 class=\"section-title\">Organizational Adoption</h2>
  <div id=\"fig_site_activity\" class=\"chart\"></div>

  <h2 class=\"section-title\">Feature Adoption Analysis</h2>
  <div id=\"fig_features\" class=\"chart\"></div>

  <h2 class=\"section-title\">Deep Dive: AskAI Usage</h2>
  <div id=\"fig_askai_sites\" class=\"chart\"></div>

  <script>
    const usageRecords = __USAGE__;
    const askaiRecords = __ASKAI__;
    const selectedSections = __SECTIONS__;

    function toDateOnly(ms) {
      const d = new Date(ms);
      return new Date(d.getFullYear(), d.getMonth(), d.getDate());
    }

    function endOfDay(date) {
      return new Date(date.getFullYear(), date.getMonth(), date.getDate(), 23, 59, 59, 999);
    }

    function mondayOfWeek(d) {
      const date = new Date(d);
      const day = (date.getDay() + 6) % 7; // Monday=0
      const monday = new Date(date);
      monday.setDate(date.getDate() - day);
      monday.setHours(0,0,0,0);
      return monday;
    }

    function filterByRange(records, start, end) {
      if (!records || records.length === 0) return [];
      const s = start ? start.getTime() : -Infinity;
      const e = end ? end.getTime() : Infinity;
      return records.filter(r => typeof r.createdAt_ms === 'number' && r.createdAt_ms >= s && r.createdAt_ms <= e);
    }

    function plotWAU(filtered) {
      const weekToAccounts = new Map();
      filtered.forEach(r => {
        const week = mondayOfWeek(new Date(r.createdAt_ms)).toISOString().slice(0,10);
        if (!weekToAccounts.has(week)) weekToAccounts.set(week, new Set());
        if (r.account) weekToAccounts.get(week).add(r.account);
      });
      const weeks = Array.from(weekToAccounts.keys()).sort();
      const wau = weeks.map(w => weekToAccounts.get(w).size);
      Plotly.newPlot('fig_wau', [{ x: weeks, y: wau, type: 'scatter', mode: 'lines+markers', line: { color: '#1f77b4' } }],
        { title: '<b>Weekly Active Users (WAU) Trend</b>', template: 'plotly_white' }, {responsive: true});
    }

    function plotHeatmap(filtered) {
      const dayOrder = ['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday'];
      const counts = Array.from({length:24}, () => Array.from({length:7}, () => 0));
      filtered.forEach(r => {
        const d = new Date(r.createdAt_ms);
        const hour = d.getHours();
        const day = (d.getDay() + 6) % 7; // Monday first
        counts[hour][day] += 1;
      });
      const zmax = counts.reduce((m, row) => Math.max(m, ...row), 0);
      Plotly.newPlot('fig_heatmap', [{
        z: counts,
        x: dayOrder,
        y: Array.from({length:24}, (_,i)=>i),
        type: 'heatmap',
        colorscale: [
          [0.0, '#eef5ff'],
          [0.2, '#d6e9ff'],
          [0.4, '#9ecae1'],
          [0.6, '#6baed6'],
          [0.8, '#3182bd'],
          [1.0, '#08519c']
        ],
        zauto: false, zmin: 0, zmax: zmax
      }], { title: '<b>User Activity Heatmap (by Day and Hour)</b>', xaxis: {title: 'Day of the Week'}, yaxis: {title: 'Hour of the Day (HKT)'}, template: 'plotly_white' }, {responsive:true});
    }

    function plotSiteActivity(filtered) {
      const map = new Map();
      filtered.forEach(r => {
        const s = r.site_code || 'Unknown';
        map.set(s, (map.get(s)||0)+1);
      });
      const labels = Array.from(map.keys());
      const values = labels.map(k => map.get(k));
      if (labels.length === 0) {
        document.getElementById('fig_site_activity').innerHTML = '<div style="text-align:center;color:#888;">No data for selected period</div>';
        return;
      }
      const total = values.length ? values.reduce((a,b)=>a+b,0) : 0;
      const ids = ['root', ...labels.map((l, i) => `site_${i}`)];
      Plotly.newPlot('fig_site_activity', [{
        type: 'treemap',
        ids: ids,
        labels: ['All Sites', ...labels],
        parents: ['', ...labels.map(_=> 'root')],
        values: [total, ...values],
        textinfo: 'label+value', branchvalues: 'total'
      }], { title: '<b>Activity Distribution by Site Code</b>', template: 'plotly_white' }, {responsive:true});
    }

    function plotFeatures(filtered) {
      const map = new Map();
      filtered.forEach(r => {
        const f = r.usage_type || 'unknown';
        map.set(f, (map.get(f)||0)+1);
      });
      const entries = Array.from(map.entries()).sort((a,b)=>a[1]-b[1]);
      const y = entries.map(e=>e[0]);
      const x = entries.map(e=>e[1]);
      Plotly.newPlot('fig_features', [{ x, y, type:'bar', orientation:'h', marker: {color:'#5a4fcf'}, text: x, textposition: 'outside' }],
        { title: '<b>What Features Are Users Exploring?</b>', xaxis: {title: 'Number of Times Feature Was Used'}, yaxis: {automargin: true}, margin: {l: 240, r: 40, t: 60, b: 60}, template: 'plotly_white' }, {responsive:true});
    }

    function plotAskAISites(filtered) {
      const map = new Map();
      filtered.forEach(r => {
        const s = r.site_code || 'Unknown';
        map.set(s, (map.get(s)||0)+1);
      });
      const entries = Array.from(map.entries()).sort((a,b)=>a[1]-b[1]);
      const y = entries.map(e=>e[0]);
      const x = entries.map(e=>e[1]);
      Plotly.newPlot('fig_askai_sites', [{ x, y, type:'bar', orientation:'h', marker: { color: x, colorscale: 'Viridis', showscale: false } }],
        { title: '<b>AskAI Pioneers: Adoption by Site</b>', xaxis: {title:'Number of AskAI Queries'}, yaxis: {title:'Site Code'}, template: 'plotly_white' }, {responsive:true});
    }

    // AskAI keywords chart removed per request

    function applyRange(start, end) {
      const usage = filterByRange(usageRecords, start, end);
      const askai = filterByRange(askaiRecords, start, end);
      if (selectedSections.includes('engagement')) {
        plotWAU(usage);
        plotHeatmap(usage);
      } else {
        document.getElementById('fig_wau').innerHTML='';
        document.getElementById('fig_heatmap').innerHTML='';
      }
      if (selectedSections.includes('adoption')) {
        plotSiteActivity(usage);
      } else {
        document.getElementById('fig_site_activity').innerHTML='';
      }
      if (selectedSections.includes('features')) {
        plotFeatures(usage);
      } else {
        document.getElementById('fig_features').innerHTML='';
      }
      if (selectedSections.includes('askai')) {
        plotAskAISites(askai);
      } else {
        document.getElementById('fig_askai_sites').innerHTML='';
      }
    }

    const btns = document.querySelectorAll('.btn-group button');
    btns.forEach(b => b.addEventListener('click', () => {
      btns.forEach(x => x.classList.remove('active'));
      b.classList.add('active');
      const val = b.dataset.range;
      const today = toDateOnly(Date.now());
      let start=null, end=null;
      if (val === 'all') { start = null; end = null; }
      else {
        const days = parseInt(val, 10);
        end = endOfDay(today);
        start = new Date(today.getTime() - (days-1)*24*3600*1000);
      }
      const sd = document.getElementById('startDate');
      const ed = document.getElementById('endDate');
      sd.value = start ? new Date(start).toISOString().slice(0,10) : '';
      ed.value = end ? new Date(end).toISOString().slice(0,10) : '';
      applyRange(start, end);
    }));

    document.getElementById('applyBtn').addEventListener('click', () => {
      const sd = document.getElementById('startDate').value;
      const ed = document.getElementById('endDate').value;
      const start = sd ? toDateOnly(Date.parse(sd)) : null;
      const end = ed ? endOfDay(toDateOnly(Date.parse(ed))) : null;
      applyRange(start, end);
    });

    (function init() {
      // Default to All time to match DOCX behavior
      document.querySelectorAll('.btn-group button').forEach(x => x.classList.remove('active'));
      const allBtn = document.querySelector('.btn-group button[data-range="all"]');
      if (allBtn) allBtn.classList.add('active');
      document.getElementById('startDate').value = '';
      document.getElementById('endDate').value = '';
      applyRange(null, null);
    })();
  </script>
</body>
</html>
"""

    html_out = html_template.replace('__GENERATED_AT__', datetime.now().strftime('%d %B %Y, %H:%M:%S HKT'))
    html_out = html_out.replace('__USAGE__', json.dumps(usage_records))
    html_out = html_out.replace('__ASKAI__', json.dumps(askai_records))
    html_out = html_out.replace('__SECTIONS__', json.dumps(selected_sections))
    # Inject KPI values
    html_out = html_out.replace('{total_accounts}', str(total_accounts))
    html_out = html_out.replace('{pro_users}', str(pro_users))
    html_out = html_out.replace('{team_users}', str(team_users))

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_out)
    print(f"  ✅ Interactive HTML saved as: {output_path}")

# --- [MODIFIED] Function to send email with embedded charts
def send_email_with_charts(figures_dict, selected_sections, total_accounts, pro_users, team_users,
                          include_word_attachment=False, word_file_path=None,
                          include_html_attachment=False, html_file_path=None):
    """Constructs and sends an email with multiple embedded Plotly charts."""

    SENDER_EMAIL = os.getenv('SENDER_EMAIL')
    SENDER_PASSWORD = os.getenv('SENDER_PASSWORD')
    RECEIVER_EMAIL = os.getenv('RECEIVER_EMAIL')

    if not all([SENDER_EMAIL, SENDER_PASSWORD, RECEIVER_EMAIL]):
        print("❌ ERROR: Missing email credentials in .env file. Cannot send email.")
        return

    print(f"  -> Preparing to send email to {RECEIVER_EMAIL}...")

    msg = MIMEMultipart('related')
    msg['Subject'] = f"KaiShing OAK Usage Report - {datetime.now().strftime('%d %B %Y')}"
    msg['From'] = SENDER_EMAIL
    msg['To'] = RECEIVER_EMAIL

    # Build HTML body based on selected sections
    html_body = f"""
    <html>
      <head>
        <style>
          body {{ font-family: Arial, sans-serif; color: #333; }}
          h1 {{ color: #1a1a1a; text-align: center; }}
          h2 {{ color: #00529B; border-bottom: 2px solid #00529B; padding-bottom: 5px; margin-top: 40px; }}
        </style>
      </head>
      <body>
        <h1>KaiShing OAK Usage Dashboard</h1>
        <p style="text-align: center; color: #666;">Report Generated: {datetime.now().strftime('%d %B %Y, %H:%M:%S HKT')}</p>
    """

    # Add overview section if selected
    if 'overview' in selected_sections:
        html_body += f"""
        <h2>Overall Platform Statistics</h2>
        <table role="presentation" border="0" cellpadding="5" cellspacing="5" width="100%">
          <tr>
            <td valign="top" width="33.33%" style="background-color:#f4f4f4; border-radius:8px; padding:20px; text-align:center;">
              <p style="font-size: 2.5em; margin: 0; color: #1a1a1a;">{total_accounts}</p>
              <p style="margin: 0; color: #333;">Total Active Accounts</p>
            </td>
            <td valign="top" width="33.33%" style="background-color:#f4f4f4; border-radius:8px; padding:20px; text-align:center;">
              <p style="font-size: 2.5em; margin: 0; color: #5c2c85;">{pro_users}</p>
              <p style="margin: 0; color: #333;">Pro Tier Users</p>
            </td>
            <td valign="top" width="33.33%" style="background-color:#f4f4f4; border-radius:8px; padding:20px; text-align:center;">
              <p style="font-size: 2.5em; margin: 0; color: #5c2c85;">{team_users}</p>
              <p style="margin: 0; color: #333;">Team Tier Users</p>
            </td>
          </tr>
        </table>
        """

    # Add engagement section if selected
    if 'engagement' in selected_sections:
        html_body += """
        <h2>Platform Engagement & Activity Trends</h2>
        <img src="cid:fig_wau" style="width:100%; max-width:800px; height:auto; display:block;"><br>
        <img src="cid:fig_heatmap" style="width:100%; max-width:800px; height:auto; display:block;"><br>
        """

    # Add adoption section if selected
    if 'adoption' in selected_sections:
        html_body += """
        <h2>Organizational Adoption</h2>
        <img src="cid:fig_site_activity" style="width:100%; max-width:800px; height:auto; display:block;"><br>
        """

    # Add features section if selected
    if 'features' in selected_sections:
        html_body += """
        <h2>Feature Adoption Analysis</h2>
        <img src="cid:fig_features" style="width:100%; max-width:800px; height:auto; display:block;"><br>
        """

    # Add AskAI section if selected
    if 'askai' in selected_sections:
        html_body += """
        <h2>Deep Dive: AskAI Usage</h2>
        <img src="cid:fig_askai_sites" style="width:100%; max-width:800px; height:auto; display:block;"><br>
        <img src="cid:fig_askai_keywords" style="width:100%; max-width:800px; height:auto; display:block;"><br>
        """

    html_body += """
        <p><i>This is an automated report.</i></p>
      </body>
    </html>
    """

    msg.attach(MIMEText(html_body, 'html'))

    # Attach each figure to the email (only if section is selected)
    figures_to_attach = {
        'fig_wau': 'engagement',
        'fig_heatmap': 'engagement',
        'fig_site_activity': 'adoption',
        'fig_features': 'features',
        'fig_askai_sites': 'askai',
        'fig_askai_keywords': 'askai'
    }

    for cid, required_section in figures_to_attach.items():
        if required_section in selected_sections and cid in figures_dict and figures_dict[cid] and figures_dict[cid].data:
            image_bytes = figures_dict[cid].to_image(format="png", width=800, height=450, scale=1)
            image = MIMEImage(image_bytes)
            image.add_header('Content-ID', f'<{cid}>')
            msg.attach(image)
            print(f"  - Attached figure '{cid}'")

    # Attach Word document if requested
    if include_word_attachment and word_file_path and os.path.exists(word_file_path):
        with open(word_file_path, "rb") as f:
            word_attachment = MIMEApplication(f.read(), _subtype="docx")
            word_attachment.add_header('Content-Disposition', 'attachment', filename=os.path.basename(word_file_path))
            msg.attach(word_attachment)
            print(f"  - Attached Word document: {os.path.basename(word_file_path)}")

    # Attach HTML report if requested
    if include_html_attachment and html_file_path and os.path.exists(html_file_path):
        with open(html_file_path, "rb") as f:
            html_attachment = MIMEApplication(f.read(), _subtype="html")
            html_attachment.add_header('Content-Disposition', 'attachment', filename=os.path.basename(html_file_path))
            msg.attach(html_attachment)
            print(f"  - Attached HTML report: {os.path.basename(html_file_path)}")

    # Send the Email
    try:
        print("  -> Connecting to SMTP server...")
        server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
        server.login(SENDER_EMAIL, SENDER_PASSWORD)
        server.send_message(msg)
        server.quit()
        print("  ✅ Email sent successfully!")
    except Exception as e:
        print(f"  ❌ ERROR sending email: {e}")

# --- [NEW] Main function with command-line interface
def main():
    parser = argparse.ArgumentParser(description='Generate KaiShing OAK Usage Report')
    parser.add_argument('--sections', nargs='+', choices=list(AVAILABLE_SECTIONS.keys()),
                       default=list(AVAILABLE_SECTIONS.keys()),
                       help='Sections to include in the report')
    parser.add_argument('--start-date', type=str,
                       help='Start date for data filtering (YYYY-MM-DD format)')
    parser.add_argument('--end-date', type=str,
                       help='End date for data filtering (YYYY-MM-DD format)')
    parser.add_argument('--export-word', action='store_true',
                       help='Export report as Word document')
    parser.add_argument('--send-email', action='store_true',
                       help='Send report via email')
    parser.add_argument('--word-file', type=str, default='oak_report.docx',
                       help='Output filename for Word document')
    parser.add_argument('--export-html', action='store_true',
                       help='Export interactive HTML report')
    parser.add_argument('--html-file', type=str, default='oak_report.html',
                       help='Output filename for HTML report')

    args = parser.parse_args()

    # Parse dates if provided
    start_date = None
    end_date = None
    if args.start_date:
        start_date = pd.to_datetime(args.start_date)
    if args.end_date:
        end_date = pd.to_datetime(args.end_date)

    # If only start date provided, set end date to today
    if start_date and not end_date:
        end_date = datetime.now()
    # If only end date provided, set start date to 30 days ago
    if end_date and not start_date:
        start_date = end_date - timedelta(days=30)

    print("--- Starting Data Fetch for Report ---")
    load_dotenv(override=True)

    # Load data
    tables_to_load = [
        "oak-account-ks", "oak-ask-ai-ks", "oak-error-log-ks", "oak-jargon-ks",
        "oak-speaker-ks", "oak-subscription-log-ks", "oak-task-ks",
        "oak-terraform-state-lock-ks", "oak-transcription-ks", "oak-usage-log-ks",
        "oak-ws-session-ks", "oak-zoom-auth-ks", "oak-zoom-integration-ks",
        "oak-zoom-session-ks"
    ]
    dataframes = {}
    for table_name in tables_to_load:
        df = get_data_from_dynamodb(
            table_name=table_name,
            aws_key=os.getenv("KAISHING_DYNAMODB_ACCESS_KEY_ID"),
            aws_secret=os.getenv("KAISHING_DYNAMODB_SECRET_ACCESS_KEY"),
            aws_region=os.getenv("KAISHING_DYNAMODB_REGION")
        )
        dataframes[table_name] = df
    print("--- Data Loading Complete ---\n")

    # Process and filter data
    print("--- Processing Data for Visualizations ---")
    users_to_exclude = ['kian.so@thinkcol.com', 'hetty.pun@thinkcol.com', 'adawan@kaishing.com.hk']
    site_code_map = {
        'eddiecheuk@kaishing.com.hk': 'HQ-IT', 'ksitsupport@kaishing.com.hk': 'HQ-IT', 'aegeancoast@kaishing.com.hk': 'AC',
        'dacychung@kaishing.com.hk': 'ICC', 'lewislam@kaishing.com.hk': 'ICC', 'Vcity@kaishing.com.hk': 'VCY',
        'yohomidtown@kaishing.com.hk': 'YMT', 'leightonhill@supreme-mgt.com.hk': 'LH', 'riva@supreme-mgt.com.hk': 'RV',
        'tpmm@kaishing.com.hk': 'TPMM', 'palmsprings@kaishing.com.hk': 'PS', 'castello@kaishing.com.hk': 'CAS',
        'newtown3@kaishing.com.hk': 'NTP3R', 'millencity@kaishing.com.hk': 'M388', 'mounthaven@kaishing.com.hk': 'MH',
        'victorwong@supreme-mgt.com.hk': 'UMA', 'epc@kaishing.com.hk': 'EPC-C', 'millencity5@kaishing.com.hk': 'MMC418',
        'apm@kaishing.com.hk': 'MMC418', 'taipocentre@kaishing.com.hk': 'TPC', 'parkisland@kaishing.com.hk': 'PI',
        'thewings3a@kaishing.com.hk': 'TW3A', 'pacificview@kaishing.com.hk': 'PV', 'cffy@chifufayuen.hk': 'CFFY',
        '98hms@kaishing.com.hk': '98HMS', 'stanford@kaishing.com.hk': 'SFV', 'lepalais@kaishing.com.hk': 'LPS',
        'avignon@kaishing.com.hk': 'AGN', 'pmt@kaishing.com.hk': 'PMT', 'mountregency@kaishing.com.hk': 'MR',
        'somerset@kaishing.com.hk': 'SOM', 'emilyho@kaishing.com.hk': 'NTPI', 'garychan@kaishing.com.hk': 'CAC',
        'dynastycourt@kaishing.com.hk': 'DC', 'eastpoint@kaishing.com.hk': 'EPCR', 'grandyoho@kaishing.com.hk': 'GYR',
        'hillsborough@kaishing.com.hk': 'HC', 'kodakhouse11@kaishing.com.hk': 'KHII', 'oceanwings@kaishing.com.hk': 'OW',
        'pokfulam@kaishing.com.hk': 'PG', 'royalpalms@kaishing.com.hk': 'RP', 'concerto@kaishing.com.hk': 'VC',
        'brownieyu@kaishing.com.hk': 'AFFC', 'hlypm@kaishing.com.hk': 'HLY', 'thewings2@kaishing.com.hk': 'TW2',
        'mayfair@kaishing.com.hk': 'MG', 'affc@kaishing.com.hk': 'AFFC', 'villabythepark@kaishing.com.hk': 'VP',
        'celestecourt@kaishing.com.hk': 'CC', 'ls@kaishing.com.hk': 'LS', 'suntuenmun@kaishing.com.hk': 'STMC',
        'lagrove@kaishing.com.hk': 'LG', 'yohowest@wespire.com.hk': 'YOW', 'yohohouse@wespire.com.hk': 'YOW',
        'kennedy38@supreme-mgt.com.hk': 'K38', 'homantinhill@supreme-mgt.com.hk': 'HMT', 'landmarkn@kaishing.com.hk': 'LN',
        'metroplaza@kaishing.com.hk': 'MP', 'yohomall-1@kaishing.com.hk': 'YM1', 'ylplaza@kaishing.com.hk': 'YLP',
        'kingspark@kaishing.com.hk': 'KPV', 'candicewong@kaishing.com.hk': 'KCC', 'rhapsody@kaishing.com.hk': 'VR',
        'lgar@kaishing.com.hk': 'LGAR', 'rseacrest@kaishing.com.hk': 'RSC', 'yukpocourt@kaishing.com.hk': 'YPC',
        'villaathena@kaishing.com.hk': 'VA', 'vincenttse@supreme-mgt.com.hk': 'VY'
    }

    account_df = dataframes['oak-account-ks'][~dataframes['oak-account-ks']['account'].isin(users_to_exclude)].copy()
    usage_log_df = dataframes['oak-usage-log-ks'][~dataframes['oak-usage-log-ks']['account'].isin(users_to_exclude)].copy()
    ask_ai_df = dataframes['oak-ask-ai-ks'][~dataframes['oak-ask-ai-ks']['account'].isin(users_to_exclude)].copy()

    # Apply time filtering if dates are provided
    if start_date and end_date:
        print(f"  -> Filtering data from {start_date.strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')}")
        usage_log_df = filter_data_by_time_period(usage_log_df, start_date, end_date)
        ask_ai_df = filter_data_by_time_period(ask_ai_df, start_date, end_date)

    for df in [usage_log_df, ask_ai_df]:
        if 'createdAt' in df.columns:
            numeric_timestamps = pd.to_numeric(df['createdAt'], errors='coerce')
            df['createdAt'] = pd.to_datetime(numeric_timestamps, unit='s', errors='coerce').dt.tz_localize('UTC')

    usage_log_valid_time_df = usage_log_df.dropna(subset=['createdAt']).copy()
    usage_log_valid_time_df['site_code'] = usage_log_valid_time_df['account'].map(site_code_map).fillna('Unknown')
    if not ask_ai_df.empty:
        ask_ai_df['site_code'] = ask_ai_df['account'].map(site_code_map).fillna('Unknown')

    # Generate figures (same as original code)
    print("--- Generating Figures for Report ---")
    total_accounts = len(account_df)
    subscription_counts = account_df['subscription_level'].value_counts(dropna=False)
    pro_users = subscription_counts.get('pro', 0)
    team_users = subscription_counts.get('team', 0)

    if not usage_log_valid_time_df.empty:
        hkt_timezone = 'Asia/Hong_Kong'
        usage_log_valid_time_df['createdAt_HKT'] = usage_log_valid_time_df['createdAt'].dt.tz_convert(hkt_timezone)
        wau_df = usage_log_valid_time_df.groupby(pd.Grouper(key='createdAt_HKT', freq='W-Mon'))['account'].nunique().reset_index()
        wau_df.rename(columns={'account': 'Weekly Active Users'}, inplace=True)
        fig_wau = px.line(wau_df, x='createdAt_HKT', y='Weekly Active Users', title='<b>Weekly Active Users (WAU) Trend</b>', labels={'createdAt_HKT': 'Week (in HKT)'}, markers=True, template='plotly_white')
        fig_wau.update_layout(title_x=0.5)
        fig_wau.update_traces(line_color='#1f77b4')
        usage_log_valid_time_df['day_of_week'] = usage_log_valid_time_df['createdAt_HKT'].dt.day_name()
        usage_log_valid_time_df['hour_of_day'] = usage_log_valid_time_df['createdAt_HKT'].dt.hour
        activity_heatmap_data = usage_log_valid_time_df.pivot_table(index='hour_of_day', columns='day_of_week', values='id', aggfunc='count').fillna(0)
        day_order = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
        activity_heatmap_data = activity_heatmap_data.reindex(columns=day_order, fill_value=0)
        fig_heatmap = go.Figure(data=go.Heatmap(z=activity_heatmap_data.values, x=activity_heatmap_data.columns, y=activity_heatmap_data.index, colorscale='Blues'))
        fig_heatmap.update_layout(title='<b>User Activity Heatmap (by Day and Hour)</b>', xaxis_title='Day of the Week', yaxis_title='Hour of the Day (HKT)', title_x=0.5)
    else:
        fig_wau = go.Figure()
        fig_heatmap = go.Figure()

    activity_by_site = usage_log_valid_time_df['site_code'].value_counts().reset_index()
    activity_by_site.columns = ['site_code', 'action_count']
    fig_site_activity = px.treemap(activity_by_site, path=[px.Constant("All Sites"), 'site_code'], values='action_count', title='<b>Activity Distribution by Site Code</b>', template='plotly_white', hover_data={'action_count': ':.0f'}, color_continuous_scale='Greens')
    fig_site_activity.update_layout(title_x=0.5, margin = dict(t=50, l=25, r=25, b=25))
    fig_site_activity.update_traces(textinfo="label+value", textfont_size=14)

    feature_usage_counts = usage_log_df['usage_type'].value_counts().reset_index()
    feature_usage_counts.columns = ['feature', 'count']
    fig_features = px.bar(feature_usage_counts, x='count', y='feature', orientation='h', title='<b>What Features Are Users Exploring?</b>', template='plotly_white', text='count')
    fig_features.update_yaxes(categoryorder='total ascending')
    fig_features.update_layout(title_x=0.5, xaxis_title='Number of Times Feature Was Used')
    fig_features.update_traces(textposition='outside', marker_color='#5a4fcf')

    if not ask_ai_df.empty:
        ask_ai_by_site = ask_ai_df['site_code'].value_counts().reset_index()
        ask_ai_by_site.columns = ['site_code', 'query_count']
        fig_askai_sites = px.bar(ask_ai_by_site, x='query_count', y='site_code', orientation='h', title='<b>AskAI Pioneers: Adoption by Site</b>', text='query_count', template='plotly_white', color='query_count', color_continuous_scale=px.colors.sequential.Viridis)
        fig_askai_sites.update_yaxes(categoryorder='total ascending')
        fig_askai_sites.update_layout(title_x=0.5, xaxis_title='Number of AskAI Queries', yaxis_title='Site Code', coloraxis_showscale=False)
        if 'question' in ask_ai_df.columns and not ask_ai_df['question'].dropna().empty:
            stopwords = set(['the', 'a', 'an', 'is', 'are', 'to', 'and', 'of', 'in', 'what', 'who', 'how', 'summary', 'summarize'])
            all_queries = ' '.join(ask_ai_df['question'].dropna().str.lower())
            words = re.findall(r'\b\w+\b', all_queries)
            filtered_words = [word for word in words if word not in stopwords and not word.isdigit()]
            top_words = Counter(filtered_words).most_common(10)
            if top_words:
                top_words_df = pd.DataFrame(top_words, columns=['keyword', 'count'])
                fig_askai_keywords = px.bar(top_words_df, x='count', y='keyword', orientation='h', title='<b>What Are Users Asking? Top AskAI Keywords</b>', template='plotly_white', color='count', color_continuous_scale=px.colors.sequential.Viridis)
                fig_askai_keywords.update_yaxes(categoryorder='total ascending')
                fig_askai_keywords.update_layout(title_x=0.5, coloraxis_showscale=False)
            else:
                fig_askai_keywords = go.Figure()
        else:
            fig_askai_keywords = go.Figure()
    else:
        fig_askai_sites = go.Figure()
        fig_askai_keywords = go.Figure()

    # Create figures dictionary
    figures_to_email = {
        'fig_wau': fig_wau,
        'fig_heatmap': fig_heatmap,
        'fig_site_activity': fig_site_activity,
        'fig_features': fig_features,
        'fig_askai_sites': fig_askai_sites,
        'fig_askai_keywords': fig_askai_keywords,
    }

    # Export Word document if requested
    if args.export_word:
        print("--- Creating Word Document ---")
        create_word_document(figures_to_email, args.sections, total_accounts, pro_users, team_users,
                           usage_log_valid_time_df, ask_ai_df, args.word_file)

    # Export interactive HTML if requested
    if args.export_html:
        print("--- Creating Interactive HTML ---")
        create_interactive_html(usage_log_df=usage_log_valid_time_df, ask_ai_df=ask_ai_df, selected_sections=args.sections,
                                total_accounts=total_accounts, pro_users=pro_users, team_users=team_users,
                                output_path=args.html_file)

    # Send email if requested
    if args.send_email:
        print("--- Sending Email ---")
        send_email_with_charts(figures_to_email, args.sections, total_accounts, pro_users, team_users,
                              include_word_attachment=args.export_word, word_file_path=args.word_file,
                              include_html_attachment=args.export_html, html_file_path=args.html_file)

    print("--- Report Generation Process Complete ---")

if __name__ == '__main__':
    main()
